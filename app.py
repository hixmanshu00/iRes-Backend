from __future__ import annotations

import io
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

import streamlit as st
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from pipeline import iter_research_pipeline
from research_runtime import ResearchRuntime

STEPS: list[str] = ["search", "reader", "writer", "critic", "refiner"]
STEP_META: dict[str, tuple[str, str, str]] = {
    "search":  ("01", "Search Agent",  "Web search via Tavily"),
    "reader":  ("02", "Reader Agent",  "Scraping top URLs"),
    "writer":  ("03", "Writer Chain",  "Drafting report"),
    "critic":  ("04", "Critic Chain",  "Scoring the report"),
    "refiner": ("05", "Refiner Chain", "Improving if score < 7.5"),
}

_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&family=Syne:wght@700;800&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --bg:#070a13; --bg2:#0b1326; --bg3:#0d152b;
  --line:#1e2d4a; --line2:#2a3f61;
  --muted:#7a8fb5; --text:#e8f0ff; --text2:#c0d0ee;
  --ca:#22d3ee; --cb:#6366f1; --cc:#34d399;
  --warn:#fbbf24; --danger:#f87171;
}
*,*::before,*::after{box-sizing:border-box;}
html,body,[class*="css"]{font-family:'Outfit',sans-serif;color:var(--text);}
.stApp{background:radial-gradient(ellipse at 4% -8%,rgba(99,102,241,.16),transparent 38%),radial-gradient(ellipse at 96% 4%,rgba(34,211,238,.12),transparent 36%),var(--bg);}
#MainMenu,header,footer{visibility:hidden;}
.block-container{max-width:1260px;padding-top:.9rem;padding-bottom:2rem;}
::-webkit-scrollbar{width:5px;}
::-webkit-scrollbar-track{background:var(--bg2);}
::-webkit-scrollbar-thumb{background:var(--line2);border-radius:3px;}

/* Sidebar */
section[data-testid="stSidebar"]{background:var(--bg2)!important;border-right:1px solid var(--line)!important;}
section[data-testid="stSidebar"]>div{padding-top:0!important;}

/* Typography */
.kicker{font-family:'JetBrains Mono',monospace;font-size:.62rem;letter-spacing:.18em;text-transform:uppercase;color:var(--ca);opacity:.9;}
.hero-title{margin:.25rem 0 .45rem;font-family:'Syne',sans-serif;font-size:clamp(2rem,4vw,2.9rem);font-weight:800;letter-spacing:-.03em;line-height:1.05;background:linear-gradient(130deg,#e8f0ff 0%,var(--ca) 45%,var(--cb) 100%);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;}
.hero-sub{color:var(--muted);font-size:.96rem;max-width:600px;line-height:1.6;margin-bottom:1.3rem;}

/* Cards */
.card{background:var(--bg2);border:1px solid var(--line);border-radius:16px;padding:1.1rem 1.2rem;}

/* Pipeline */
.pipeline-wrap{display:flex;flex-direction:column;}
.pipe-step{display:flex;align-items:flex-start;gap:.7rem;padding:.6rem 0;position:relative;}
.pipe-step:not(:last-child)::after{content:'';position:absolute;left:15px;top:36px;bottom:-10px;width:2px;background:var(--line);}
.pipe-step.done::after{background:rgba(52,211,153,.25);}
.pipe-step.run::after{background:rgba(34,211,238,.25);}
.pipe-icon{width:32px;height:32px;border-radius:50%;flex-shrink:0;display:flex;align-items:center;justify-content:center;font-family:'JetBrains Mono',monospace;font-size:.72rem;font-weight:700;background:var(--bg3);border:2px solid var(--line);color:var(--muted);}
.pipe-step.run  .pipe-icon{border-color:var(--ca);background:rgba(34,211,238,.1);color:var(--ca);animation:pulse 1.6s ease-in-out infinite;}
.pipe-step.done .pipe-icon{border-color:var(--cc);background:rgba(52,211,153,.12);color:var(--cc);}
.pipe-step.err  .pipe-icon{border-color:var(--danger);background:rgba(248,113,113,.1);color:var(--danger);}
@keyframes pulse{0%,100%{box-shadow:0 0 10px rgba(34,211,238,.2);}50%{box-shadow:0 0 20px rgba(34,211,238,.45);}}
.pipe-body{flex:1;min-width:0;}
.pipe-label{font-size:.87rem;font-weight:600;color:var(--text2);line-height:1.2;}
.pipe-step.run  .pipe-label{color:var(--ca);}
.pipe-step.done .pipe-label{color:var(--text);}
.pipe-meta{font-family:'JetBrains Mono',monospace;font-size:.66rem;color:var(--muted);margin-top:.1rem;}

/* Badges */
.badge{display:inline-block;padding:.1rem .45rem;border-radius:999px;font-family:'JetBrains Mono',monospace;font-size:.6rem;font-weight:600;letter-spacing:.06em;text-transform:uppercase;}
.badge-wait{background:rgba(122,143,181,.08);color:var(--muted);border:1px solid var(--line);}
.badge-run{background:rgba(34,211,238,.1);color:var(--ca);border:1px solid rgba(34,211,238,.3);}
.badge-done{background:rgba(52,211,153,.1);color:var(--cc);border:1px solid rgba(52,211,153,.3);}
.badge-err{background:rgba(248,113,113,.1);color:var(--danger);border:1px solid rgba(248,113,113,.3);}

/* Metrics */
.metrics-grid{display:grid;grid-template-columns:repeat(5,1fr);gap:.55rem;margin:.4rem 0;}
.metric-tile{background:var(--bg3);border:1px solid var(--line);border-radius:11px;padding:.65rem .75rem;text-align:center;}
.metric-val{font-family:'JetBrains Mono',monospace;font-size:1.15rem;font-weight:500;color:var(--ca);line-height:1;}
.metric-label{font-size:.62rem;color:var(--muted);margin-top:.2rem;text-transform:uppercase;letter-spacing:.07em;}

/* Sources */
.source-card{background:var(--bg3);border:1px solid var(--line);border-radius:12px;padding:.8rem 1rem;margin-bottom:.55rem;}
.source-title{font-weight:600;font-size:.9rem;color:var(--text);}
.source-url{font-family:'JetBrains Mono',monospace;font-size:.67rem;color:var(--ca);margin:.15rem 0 .3rem;word-break:break-all;}
.source-snip{font-size:.8rem;color:var(--muted);line-height:1.55;}
.source-meta{display:flex;align-items:center;gap:.45rem;margin-top:.5rem;}
.conf-wrap{flex:1;height:3px;border-radius:99px;background:var(--line);overflow:hidden;}
.conf-fill{height:100%;border-radius:99px;}
.tag{font-family:'JetBrains Mono',monospace;font-size:.59rem;color:var(--muted);background:var(--bg2);border:1px solid var(--line);border-radius:5px;padding:.1rem .3rem;white-space:nowrap;}

/* Report + Feedback */
.report-box{background:var(--bg3);border:1px solid var(--line);border-radius:14px;padding:1.2rem 1.5rem;line-height:1.8;font-size:.93rem;color:var(--text2);}
.feedback-box{background:var(--bg3);border:1px solid var(--line);border-radius:14px;padding:1rem 1.2rem;}

/* Score circle */
.score-circle{width:66px;height:66px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-family:'Syne',sans-serif;font-size:1.3rem;font-weight:800;flex-shrink:0;}
.score-hi{background:rgba(52,211,153,.12);border:2.5px solid var(--cc);color:var(--cc);}
.score-mid{background:rgba(251,191,36,.1);border:2.5px solid var(--warn);color:var(--warn);}
.score-lo{background:rgba(248,113,113,.1);border:2.5px solid var(--danger);color:var(--danger);}

/* Section divider */
.sec-hdr{display:flex;align-items:center;gap:.6rem;margin:1.5rem 0 .7rem;}
.sec-lbl{font-family:'JetBrains Mono',monospace;font-size:.63rem;color:var(--muted);text-transform:uppercase;letter-spacing:.12em;white-space:nowrap;}
.sec-sub{font-size:.67rem;color:var(--muted);white-space:nowrap;}
.sec-line{flex:1;height:1px;background:var(--line);}

/* Inputs */
.stTextInput>div>div>input{background:var(--bg3)!important;border:1px solid var(--line2)!important;border-radius:12px!important;color:var(--text)!important;font-family:'Outfit',sans-serif!important;font-size:.93rem!important;}
.stTextInput>div>div>input:focus{border-color:var(--ca)!important;box-shadow:0 0 0 3px rgba(34,211,238,.08)!important;}

/* Main button */
.stButton>button{border:none!important;border-radius:12px!important;font-family:'Outfit',sans-serif!important;font-weight:600!important;font-size:.92rem!important;background:linear-gradient(90deg,var(--ca),#60a5fa,var(--cb))!important;color:#031021!important;}
.stButton>button:hover{filter:brightness(1.07)!important;}

/* Sidebar buttons — override gradient */
section[data-testid="stSidebar"] .stButton>button{background:transparent!important;border:1px solid var(--line)!important;color:var(--text2)!important;font-size:.82rem!important;text-align:left!important;filter:none!important;}
section[data-testid="stSidebar"] .stButton>button:hover{border-color:var(--ca)!important;color:var(--ca)!important;}

/* Download buttons */
.stDownloadButton>button{background:var(--bg3)!important;border:1px solid var(--line2)!important;border-radius:10px!important;color:var(--text2)!important;font-family:'JetBrains Mono',monospace!important;font-size:.74rem!important;}
.stDownloadButton>button:hover{border-color:var(--ca)!important;color:var(--ca)!important;}
</style>
"""

_BRAND = """
<div style="padding:1.3rem 1rem .8rem;">
  <div style="font-family:'JetBrains Mono',monospace;font-size:.59rem;color:#7a8fb5;letter-spacing:.16em;text-transform:uppercase;margin-bottom:.2rem;">Research Studio</div>
  <div style="font-family:'Syne',sans-serif;font-size:1.45rem;font-weight:800;background:linear-gradient(130deg,#22d3ee,#6366f1);-webkit-background-clip:text;-webkit-text-fill-color:transparent;">iRes</div>
</div>
<hr style="border:none;border-top:1px solid #1e2d4a;margin:.1rem 0 .75rem;">
"""


# ── State ─────────────────────────────────────────────────────────────────────

def _init() -> None:
    for k, v in {
        "topic": "", "results": {}, "view": "new",
        "loaded_run_id": None, "export_bytes": {},
    }.items():
        if k not in st.session_state:
            st.session_state[k] = v


def _rt() -> ResearchRuntime:
    if "rt" not in st.session_state:
        st.session_state.rt = ResearchRuntime()
    return st.session_state.rt


# ── Sidebar ───────────────────────────────────────────────────────────────────

def _sidebar() -> None:
    runs = _rt().list_runs(limit=25)
    with st.sidebar:
        st.markdown(_BRAND, unsafe_allow_html=True)

        if st.button("＋  New Research", use_container_width=True, key="new_btn"):
            st.session_state.update(results={}, export_bytes={}, view="new", loaded_run_id=None)
            st.rerun()

        if not runs:
            st.markdown(
                "<div style='padding:.9rem 1rem;color:#7a8fb5;font-size:.8rem;text-align:center;'>"
                "No runs yet.<br>Run your first pipeline above.</div>",
                unsafe_allow_html=True)
            return

        st.markdown(
            "<div style='padding:.7rem .1rem .2rem;font-family:JetBrains Mono,monospace;"
            "font-size:.59rem;color:#7a8fb5;text-transform:uppercase;letter-spacing:.14em;'>Recent Runs</div>",
            unsafe_allow_html=True)

        for run in runs:
            score = f"  {run['critic_score']:.1f}/10" if run["critic_score"] else ""
            label = run["topic"]
            short = (label[:26] + "…") if len(label) > 26 else label
            if st.button(short + score, key=f"h_{run['run_id']}", use_container_width=True):
                loaded = _rt().load_run(run["run_id"])
                if loaded:
                    st.session_state.update(
                        results=loaded, topic=loaded["topic"],
                        loaded_run_id=run["run_id"], view="loaded", export_bytes={})
                    st.rerun()


# ── Pipeline HTML ──────────────────────────────────────────────────────────────

def _step_html(sid: str, status: str, lat: float | None) -> str:
    num, label, desc = STEP_META[sid]
    if status == "done":
        ic, ch = "done", "✓"
        badge = "<span class='badge badge-done'>Done</span>"
        lat_s = f"&nbsp;<span style='color:#34d399'>{lat:.1f}s</span>" if lat else ""
    elif status == "running":
        ic, ch = "run", num
        badge = "<span class='badge badge-run'>Running</span>"
        lat_s = ""
    elif status == "error":
        ic, ch = "err", "✕"
        badge = "<span class='badge badge-err'>Error</span>"
        lat_s = ""
    else:
        ic, ch = "wait", num
        badge = "<span class='badge badge-wait'>Waiting</span>"
        lat_s = ""
    return (f"<div class='pipe-step {ic}'>"
            f"<div class='pipe-icon'>{ch}</div>"
            f"<div class='pipe-body'>"
            f"<div class='pipe-label'>{label}&nbsp;{badge}</div>"
            f"<div class='pipe-meta'>{desc}{lat_s}</div>"
            f"</div></div>")


def _pipeline_card(done: set[str], running: str | None, metrics: dict, errors: dict) -> str:
    inner = ""
    for sid in STEPS:
        if sid in done:
            status = "error" if errors.get(sid) else "done"
            lat = metrics.get(f"{sid}_latency_s")
        elif sid == running:
            status, lat = "running", None
        else:
            status, lat = "wait", None
        inner += _step_html(sid, status, lat)
    return (f"<div class='card'>"
            f"<div class='kicker' style='margin-bottom:.6rem;'>Live Orchestration</div>"
            f"<div class='pipeline-wrap'>{inner}</div>"
            f"</div>")


def _metrics_card(metrics: dict) -> str:
    tiles = ""
    for sid in STEPS:
        v = metrics.get(f"{sid}_latency_s")
        vs = f"{v:.1f}s" if v else "—"
        tiles += (f"<div class='metric-tile'>"
                  f"<div class='metric-val'>{vs}</div>"
                  f"<div class='metric-label'>{STEP_META[sid][1]}</div>"
                  f"</div>")
    total = metrics.get("total_latency_s", 0)
    ts = f"{total:.1f}s" if total else "—"
    fb = [k.replace("_fallback_used", "").title() for k, v in metrics.items() if "_fallback_used" in k and v]
    fbn = f" · Fallback: {', '.join(fb)}" if fb else ""
    return (f"<div class='card'>"
            f"<div class='kicker' style='margin-bottom:.6rem;'>Step Metrics</div>"
            f"<div class='metrics-grid'>{tiles}</div>"
            f"<div style='text-align:center;font-family:JetBrains Mono,monospace;"
            f"font-size:.64rem;color:#7a8fb5;margin-top:.3rem;'>"
            f"Total&nbsp;<span style='color:#22d3ee;'>{ts}</span>{fbn}</div>"
            f"</div>")


# ── Helpers ────────────────────────────────────────────────────────────────────

def _divider(label: str, sub: str = "") -> None:
    sub_h = f"<span class='sec-sub'>{sub}</span>" if sub else ""
    st.markdown(
        f"<div class='sec-hdr'><span class='sec-lbl'>{label}</span>"
        f"{sub_h}<div class='sec-line'></div></div>",
        unsafe_allow_html=True)


def _conf_color(c: float) -> str:
    return "#34d399" if c >= .7 else "#fbbf24" if c >= .5 else "#f87171"


# ── Results ────────────────────────────────────────────────────────────────────

def _render_results(r: dict, topic: str) -> None:
    metrics = r.get("metrics", {})
    errors  = r.get("errors",  {})

    # Pipeline overview + metrics
    lc, rc = st.columns([1, 1.6], gap="large")
    with lc:
        st.markdown(_pipeline_card(set(STEPS), None, metrics, errors), unsafe_allow_html=True)
    with rc:
        if metrics:
            st.markdown(_metrics_card(metrics), unsafe_allow_html=True)

    # ── Primary report (refined if available, else original) ──────────────────
    refined   = r.get("refined_report", "")
    was_refined = r.get("was_refined", False)
    display_report = refined if (refined and was_refined) else r.get("report", "")

    if display_report:
        if was_refined and refined:
            _divider("Research Report",
                     "<span style='background:rgba(52,211,153,.15);color:#34d399;"
                     "border:1px solid rgba(52,211,153,.35);border-radius:6px;"
                     "padding:.1rem .5rem;font-size:.65rem;font-family:JetBrains Mono,monospace;"
                     "letter-spacing:.06em;text-transform:uppercase;margin-left:.4rem;'>AI-Refined</span>")
        else:
            _divider("Research Report")
        st.markdown(f"<div class='report-box'>{display_report}</div>", unsafe_allow_html=True)

        # Show original draft inside expander if refinement happened
        if was_refined and r.get("report"):
            with st.expander("View original draft (before refinement)"):
                st.markdown(f"<div class='report-box' style='opacity:.75;'>{r['report']}</div>",
                            unsafe_allow_html=True)

    # ── Critic feedback ───────────────────────────────────────────────────────
    if r.get("feedback"):
        fb   = r["feedback"]
        m    = re.search(r"Score:\s*(\d+(?:\.\d+)?)\s*/\s*10", fb, re.I)
        sc   = float(m.group(1)) if m else None
        sc_s = f"{sc:.1f}" if sc is not None else "?"
        sc_cls = "score-hi" if (sc or 0) >= 7.5 else "score-mid" if (sc or 0) >= 5 else "score-lo"
        _divider("Critic Feedback", f"Score {sc_s}/10" if sc else "")
        c1, c2 = st.columns([1, 8])
        with c1:
            st.markdown(f"<div class='score-circle {sc_cls}'>{sc_s}</div>", unsafe_allow_html=True)
        with c2:
            st.markdown(
                f"<div class='feedback-box'><pre style='white-space:pre-wrap;font-family:inherit;"
                f"font-size:.88rem;margin:0;color:#c0d0ee;'>{fb}</pre></div>",
                unsafe_allow_html=True)

    # ── Sources (collapsed by default) ────────────────────────────────────────
    sources = r.get("sources", [])
    if sources:
        with st.expander(f"Sources ({len(sources)} discovered)", expanded=False):
            for src in sources:
                conf  = src.get("confidence", .5)
                color = _conf_color(conf)
                pct   = int(conf * 100)
                st.markdown(
                    f"<div class='source-card'>"
                    f"<div class='source-title'>{src.get('title','Untitled')}</div>"
                    f"<div class='source-url'>{src.get('url','')}</div>"
                    f"<div class='source-snip'>{src.get('snippet','')[:220]}</div>"
                    f"<div class='source-meta'>"
                    f"<span class='tag'>{src.get('domain','Unknown')}</span>"
                    f"<span class='tag'>{src.get('publish_date','Unknown')}</span>"
                    f"<div class='conf-wrap'><div class='conf-fill' style='width:{pct}%;background:{color};'></div></div>"
                    f"<span style='font-family:JetBrains Mono,monospace;font-size:.63rem;color:{color};'>{pct}%</span>"
                    f"</div></div>",
                    unsafe_allow_html=True)

    # ── Exports ───────────────────────────────────────────────────────────────
    if display_report:
        _divider("Export Bundle")
        if st.button("Generate Exports  (.md  ·  .json  ·  .docx  ·  .pdf)"):
            st.session_state.export_bytes = _make_exports(topic, r)
        if st.session_state.get("export_bytes"):
            mimes = {
                ".md":   "text/markdown",
                ".json": "application/json",
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".pdf":  "application/pdf",
            }
            cols = st.columns(4)
            for i, (fname, data) in enumerate(st.session_state.export_bytes.items()):
                ext = Path(fname).suffix.lower()
                with cols[i % 4]:
                    st.download_button(
                        f"↓ {ext.upper()}", data=data, file_name=fname,
                        mime=mimes.get(ext, ""), use_container_width=True)


def _make_exports(topic: str, r: dict) -> dict[str, bytes]:
    ts     = datetime.now().strftime("%Y%m%d_%H%M%S")
    base   = f"ires_{ts}"
    refined = r.get("refined_report", "")
    was_refined = r.get("was_refined", False)
    report = refined if (refined and was_refined) else r.get("report", "")
    fb     = r.get("feedback", "")
    srcs   = r.get("sources", [])

    j_bytes = json.dumps(
        {"topic": topic, "report": report, "feedback": fb, "sources": srcs}, indent=2
    ).encode()

    doc = Document()
    doc.add_heading(topic, level=1)
    for p in report.split("\n\n"):
        doc.add_paragraph(p.strip())
    doc.add_page_break()
    doc.add_heading("Critic Feedback", level=2)
    doc.add_paragraph(fb)
    dbuf = io.BytesIO()
    doc.save(dbuf)

    pbuf = io.BytesIO()
    cv   = canvas.Canvas(pbuf, pagesize=A4)
    _, h = A4
    y = h - 40
    cv.setFont("Helvetica-Bold", 13)
    cv.drawString(35, y, topic[:95])
    y -= 20
    cv.setFont("Helvetica", 9)
    for line in (report + "\n\nCritic Feedback:\n" + fb).splitlines():
        if y < 55:
            cv.showPage()
            cv.setFont("Helvetica", 9)
            y = h - 40
        cv.drawString(35, y, line[:118])
        y -= 12
    cv.save()

    return {
        f"{base}.md":   report.encode(),
        f"{base}.json": j_bytes,
        f"{base}.docx": dbuf.getvalue(),
        f"{base}.pdf":  pbuf.getvalue(),
    }


# ── Main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    st.set_page_config(
        page_title="iRes Studio", page_icon="🔬",
        layout="wide", initial_sidebar_state="expanded")
    st.markdown(_CSS, unsafe_allow_html=True)
    _init()
    _sidebar()

    st.markdown("""
<div style="margin-bottom:1.5rem;">
  <div class="kicker">Multi-Agent · Autonomous Research</div>
  <h1 class="hero-title">iRes Studio</h1>
  <p class="hero-sub">Search → Scrape → Write → Critique → Refine. Powered by Mistral + Tavily.</p>
</div>""", unsafe_allow_html=True)

    ic, rc = st.columns([5, 1])
    with ic:
        topic = st.text_input(
            "topic", value=st.session_state.topic, label_visibility="collapsed",
            placeholder="e.g.  Impact of AI agents on software engineering in 2026") or ""
        st.session_state.topic = topic
    with rc:
        run = st.button("Run Pipeline", use_container_width=True)

    st.markdown("<div style='height:.2rem'></div>", unsafe_allow_html=True)

    if run and not topic.strip():
        st.warning("Enter a research topic to begin.")
        return

    # ── Execute ───────────────────────────────────────────────────────────────
    if run and topic.strip():
        st.session_state.update(results={}, export_bytes={}, view="new", loaded_run_id=None)

        pc, mc = st.columns([1, 1.6], gap="large")
        with pc:
            pipe_ph = st.empty()
        with mc:
            met_ph = st.empty()

        done: set[str] = set()
        final: dict[str, Any] = {}

        pipe_ph.markdown(_pipeline_card(done, "search", {}, {}), unsafe_allow_html=True)

        for sid, state in iter_research_pipeline(topic.strip()):
            done.add(sid)
            met  = state.get("metrics", {})
            errs = state.get("errors",  {})
            idx  = STEPS.index(sid)
            nxt  = STEPS[idx + 1] if idx + 1 < len(STEPS) else None

            pipe_ph.markdown(_pipeline_card(done, nxt, met, errs), unsafe_allow_html=True)
            met_ph.markdown(_metrics_card(met), unsafe_allow_html=True)
            final = state

        st.session_state.results = final
        st.session_state.view    = "completed"
        st.rerun()

    # ── Show results ──────────────────────────────────────────────────────────
    r = st.session_state.results
    if r:
        _render_results(r, st.session_state.topic or "")


if __name__ == "__main__":
    main()
