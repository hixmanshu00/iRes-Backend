from __future__ import annotations

import hashlib
import json
import os
import re
import sqlite3
import time
import uuid
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable
from urllib.parse import urlparse

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from tenacity import retry, stop_after_attempt, wait_exponential

from agents import (
    build_search_agent,
    build_url_scraper_agent,
    critic_chain,
    fallback_critic_chain,
    fallback_refiner_chain,
    fallback_writer_chain,
    get_model_params,
    refiner_chain,
    writer_chain,
)


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def score_from_feedback(feedback: str) -> float | None:
    m = re.search(r"Score:\s*(\d+(?:\.\d+)?)\s*/\s*10", feedback, flags=re.I)
    if not m:
        return None
    try:
        raw = float(m.group(1))
        return max(0.0, min(10.0, raw))
    except ValueError:
        return None


def parse_sources(search_text: str) -> list[dict[str, Any]]:
    blocks = [b.strip() for b in re.split(r"\n\s*\n", search_text) if b.strip()]
    sources: list[dict[str, Any]] = []
    for idx, block in enumerate(blocks, start=1):
        title = ""
        url = ""
        snippet = ""
        for line in block.splitlines():
            lower = line.lower().strip()
            if lower.startswith("title:"):
                title = line.split(":", 1)[1].strip()
            elif lower.startswith("url:"):
                url = line.split(":", 1)[1].strip()
            elif lower.startswith("snippet:"):
                snippet = line.split(":", 1)[1].strip()
        if not url:
            match = re.search(r"https?://[^\s\]]+", block)
            if match:
                url = match.group(0)
        domain = urlparse(url).netloc if url else ""
        pub_match = re.search(r"(20\d{2}-\d{2}-\d{2})", block)
        publish_date = pub_match.group(1) if pub_match else "Unknown"
        confidence = 0.35
        if "arxiv" in domain or ".gov" in domain or ".edu" in domain:
            confidence += 0.3
        if publish_date != "Unknown":
            confidence += 0.2
        if len(snippet) > 120:
            confidence += 0.15
        confidence = min(0.98, confidence)
        sources.append(
            {
                "id": f"src_{idx}",
                "title": title or f"Source {idx}",
                "url": url,
                "domain": domain or "Unknown",
                "publish_date": publish_date,
                "snippet": snippet or block[:280],
                "confidence": round(confidence, 2),
                "approved": True,
            }
        )
    return sources


@dataclass
class StepResult:
    output: str
    latency_s: float
    from_cache: bool
    attempts: int
    fallback_used: bool
    error: str | None = None


class ResearchRuntime:
    def __init__(self, db_path: str = "research_runs.db", log_path: str = "observability.log.jsonl") -> None:
        self.db_path = db_path
        self.log_path = log_path
        self._init_db()

    def _conn(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        return conn

    def _init_db(self) -> None:
        with self._conn() as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS run_history (
                    run_id TEXT PRIMARY KEY,
                    topic TEXT NOT NULL,
                    created_at TEXT NOT NULL,
                    model_params TEXT NOT NULL,
                    sources_json TEXT,
                    approved_source_ids_json TEXT,
                    search_results TEXT,
                    scraped_content TEXT,
                    report TEXT,
                    refined_report TEXT,
                    feedback TEXT,
                    critic_score REAL,
                    metrics_json TEXT,
                    errors_json TEXT
                )
                """
            )
            # migration: add refined_report to existing databases
            try:
                conn.execute("ALTER TABLE run_history ADD COLUMN refined_report TEXT")
            except Exception:
                pass
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS step_cache (
                    cache_key TEXT PRIMARY KEY,
                    step_name TEXT NOT NULL,
                    payload TEXT NOT NULL,
                    created_at TEXT NOT NULL
                )
                """
            )
            conn.commit()

    def _log_event(self, event: dict[str, Any]) -> None:
        record = {"timestamp": utc_now_iso(), **event}
        with open(self.log_path, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=True) + "\n")

    def _cache_key(self, step_name: str, payload: str) -> str:
        return hashlib.sha256(f"{step_name}:{payload}".encode("utf-8")).hexdigest()

    def _get_cache(self, step_name: str, payload: str) -> str | None:
        key = self._cache_key(step_name, payload)
        with self._conn() as conn:
            row = conn.execute(
                "SELECT payload FROM step_cache WHERE cache_key = ? AND step_name = ?",
                (key, step_name),
            ).fetchone()
        return row["payload"] if row else None

    def _set_cache(self, step_name: str, payload_input: str, payload_output: str) -> None:
        key = self._cache_key(step_name, payload_input)
        if not isinstance(payload_output, str):
            try:
                payload_output = json.dumps(payload_output, ensure_ascii=True)
            except Exception:  # noqa: BLE001
                payload_output = str(payload_output)
        with self._conn() as conn:
            conn.execute(
                """
                INSERT OR REPLACE INTO step_cache (cache_key, step_name, payload, created_at)
                VALUES (?, ?, ?, ?)
                """,
                (key, step_name, payload_output, utc_now_iso()),
            )
            conn.commit()

    def _call_with_timeout(self, fn: Callable[[], str], timeout_s: int) -> str:
        with ThreadPoolExecutor(max_workers=1) as ex:
            future = ex.submit(fn)
            try:
                return future.result(timeout=timeout_s)
            except FuturesTimeoutError as exc:
                raise TimeoutError(f"Step timed out after {timeout_s}s") from exc

    def _extract_agent_output(self, response: Any) -> str:
        if isinstance(response, dict) and "messages" in response:
            content = response["messages"][-1].content
            if isinstance(content, list):
                parts = []
                for block in content:
                    if isinstance(block, dict):
                        parts.append(block.get("text", ""))
                    elif isinstance(block, str):
                        parts.append(block)
                return "\n".join(parts)
            return str(content)
        return str(response)

    def _run_step(
        self,
        run_id: str,
        step_name: str,
        input_payload: str,
        fn: Callable[[], str],
        timeout_s: int,
        use_cache: bool = True,
    ) -> StepResult:
        started = time.perf_counter()
        cached = self._get_cache(step_name, input_payload) if use_cache else None
        if cached is not None:
            latency = time.perf_counter() - started
            self._log_event(
                {
                    "run_id": run_id,
                    "step": step_name,
                    "status": "ok",
                    "cached": True,
                    "latency_s": latency,
                    "attempts": 0,
                }
            )
            return StepResult(cached, latency, True, 0, False)

        attempt_count = {"n": 0}

        @retry(wait=wait_exponential(multiplier=1, min=1, max=8), stop=stop_after_attempt(3), reraise=True)
        def _attempt() -> str:
            attempt_count["n"] += 1
            return self._call_with_timeout(fn, timeout_s=timeout_s)

        try:
            out = _attempt()
            out = out or ""
            self._set_cache(step_name, input_payload, out)
            latency = time.perf_counter() - started
            self._log_event(
                {
                    "run_id": run_id,
                    "step": step_name,
                    "status": "ok",
                    "cached": False,
                    "latency_s": latency,
                    "attempts": attempt_count["n"],
                }
            )
            return StepResult(out, latency, False, attempt_count["n"], False)
        except Exception as exc:  # noqa: BLE001
            latency = time.perf_counter() - started
            self._log_event(
                {
                    "run_id": run_id,
                    "step": step_name,
                    "status": "error",
                    "cached": False,
                    "latency_s": latency,
                    "attempts": attempt_count["n"],
                    "error": str(exc),
                }
            )
            return StepResult("", latency, False, attempt_count["n"], False, str(exc))

    def run_search(self, run_id: str, topic: str, timeout_s: int = 35) -> StepResult:
        search_agent = build_search_agent()
        payload = topic.strip()
        today = datetime.now()
        date_str = today.strftime("%B %d, %Y")
        year = today.year
        prompt = (
            f"Today's date is {date_str}. You MUST only use information from your search tool results — "
            f"do NOT answer from training knowledge or memory.\n\n"
            f"Search for the most recent ({year}) information about: {topic}"
        )
        return self._run_step(
            run_id=run_id,
            step_name="search",
            input_payload=payload,
            timeout_s=timeout_s,
            fn=lambda: self._extract_agent_output(
                search_agent.invoke({"messages": [("user", prompt)]})
            ),
            use_cache=False,
        )

    def run_reader(self, run_id: str, topic: str, search_results: str, timeout_s: int = 40) -> StepResult:
        urls = re.findall(r"https?://[^\s\)\]\"']+", search_results)
        # deduplicate while preserving order, take top 3
        seen: set[str] = set()
        top_urls: list[str] = []
        for u in urls:
            if u not in seen:
                seen.add(u)
                top_urls.append(u)
            if len(top_urls) == 3:
                break

        payload = f"{topic.strip()}::{'|'.join(top_urls)}"

        def _scrape_all() -> str:
            import requests
            from bs4 import BeautifulSoup

            def _fetch(url: str) -> str:
                try:
                    r = requests.get(url, timeout=8, headers={"User-Agent": "Mozilla/5.0"})
                    r.raise_for_status()
                    soup = BeautifulSoup(r.text, "html.parser")
                    for tag in soup(["script", "style", "noscript", "nav", "header", "footer", "aside"]):
                        tag.decompose()
                    text = soup.get_text(separator="\n", strip=True)
                    return f"[{url}]\n{text[:3000]}"
                except Exception as exc:
                    return f"[{url}] FAILED: {exc}"

            if not top_urls:
                return "No URLs found in search results."

            with ThreadPoolExecutor(max_workers=3) as ex:
                results = list(ex.map(_fetch, top_urls))

            return "\n\n---\n\n".join(results)

        return self._run_step(
            run_id=run_id,
            step_name="reader",
            input_payload=payload,
            timeout_s=timeout_s,
            fn=_scrape_all,
            use_cache=True,
        )

    def run_writer(self, run_id: str, topic: str, research_text: str, timeout_s: int = 55) -> StepResult:
        payload = f"{topic.strip()}::{hashlib.sha256(research_text.encode('utf-8')).hexdigest()}"

        primary = self._run_step(
            run_id=run_id,
            step_name="writer",
            input_payload=payload,
            timeout_s=timeout_s,
            fn=lambda: writer_chain.invoke({"topic": topic, "research": research_text}),
            use_cache=False,
        )
        if not primary.error:
            return primary

        started = time.perf_counter()
        try:
            out = self._call_with_timeout(
                lambda: fallback_writer_chain.invoke({"topic": topic, "research": research_text}),
                timeout_s=timeout_s,
            )
            latency = time.perf_counter() - started
            self._log_event(
                {
                    "run_id": run_id,
                    "step": "writer_fallback",
                    "status": "ok",
                    "latency_s": latency,
                    "attempts": 1,
                }
            )
            return StepResult(out, primary.latency_s + latency, False, primary.attempts + 1, True)
        except Exception as exc:  # noqa: BLE001
            self._log_event(
                {
                    "run_id": run_id,
                    "step": "writer_fallback",
                    "status": "error",
                    "latency_s": time.perf_counter() - started,
                    "attempts": 1,
                    "error": str(exc),
                }
            )
            return StepResult("", primary.latency_s, False, primary.attempts + 1, True, str(exc))

    def run_critic(self, run_id: str, report: str, timeout_s: int = 40) -> StepResult:
        primary = self._run_step(
            run_id=run_id,
            step_name="critic",
            input_payload=hashlib.sha256(report.encode("utf-8")).hexdigest(),
            timeout_s=timeout_s,
            fn=lambda: critic_chain.invoke({"report": report}),
            use_cache=False,
        )
        if not primary.error:
            return primary
        started = time.perf_counter()
        try:
            out = self._call_with_timeout(lambda: fallback_critic_chain.invoke({"report": report}), timeout_s=timeout_s)
            latency = time.perf_counter() - started
            self._log_event(
                {"run_id": run_id, "step": "critic_fallback", "status": "ok", "latency_s": latency, "attempts": 1}
            )
            return StepResult(out, primary.latency_s + latency, False, primary.attempts + 1, True)
        except Exception as exc:  # noqa: BLE001
            self._log_event(
                {
                    "run_id": run_id,
                    "step": "critic_fallback",
                    "status": "error",
                    "latency_s": time.perf_counter() - started,
                    "attempts": 1,
                    "error": str(exc),
                }
            )
            return StepResult("", primary.latency_s, False, primary.attempts + 1, True, str(exc))

    def run_refiner(
        self,
        run_id: str,
        topic: str,
        report: str,
        feedback: str,
        score: float | None,
        timeout_s: int = 60,
        score_threshold: float = 7.5,
    ) -> StepResult:
        if score is not None and score >= score_threshold:
            return StepResult(report, 0.0, False, 0, False)

        primary = self._run_step(
            run_id=run_id,
            step_name="refiner",
            input_payload=hashlib.sha256((report + feedback).encode("utf-8")).hexdigest(),
            timeout_s=timeout_s,
            fn=lambda: refiner_chain.invoke({"topic": topic, "report": report, "feedback": feedback}),
            use_cache=False,
        )
        if not primary.error:
            return primary

        started = time.perf_counter()
        try:
            out = self._call_with_timeout(
                lambda: fallback_refiner_chain.invoke({"topic": topic, "report": report, "feedback": feedback}),
                timeout_s=timeout_s,
            )
            latency = time.perf_counter() - started
            self._log_event({"run_id": run_id, "step": "refiner_fallback", "status": "ok", "latency_s": latency, "attempts": 1})
            return StepResult(out, primary.latency_s + latency, False, primary.attempts + 1, True)
        except Exception as exc:  # noqa: BLE001
            self._log_event(
                {"run_id": run_id, "step": "refiner_fallback", "status": "error",
                 "latency_s": time.perf_counter() - started, "attempts": 1, "error": str(exc)}
            )
            return StepResult(report, primary.latency_s, False, primary.attempts + 1, True, str(exc))

    def save_run(
        self,
        run_id: str,
        topic: str,
        sources: list[dict[str, Any]],
        approved_source_ids: list[str],
        search_results: str,
        scraped_content: str,
        report: str,
        refined_report: str,
        feedback: str,
        metrics: dict[str, Any],
        errors: dict[str, str],
    ) -> None:
        with self._conn() as conn:
            conn.execute(
                """
                INSERT OR REPLACE INTO run_history (
                    run_id, topic, created_at, model_params, sources_json, approved_source_ids_json,
                    search_results, scraped_content, report, refined_report, feedback,
                    critic_score, metrics_json, errors_json
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    run_id,
                    topic,
                    utc_now_iso(),
                    json.dumps(get_model_params(), ensure_ascii=True),
                    json.dumps(sources, ensure_ascii=True),
                    json.dumps(approved_source_ids, ensure_ascii=True),
                    search_results,
                    scraped_content,
                    report,
                    refined_report,
                    feedback,
                    score_from_feedback(feedback),
                    json.dumps(metrics, ensure_ascii=True),
                    json.dumps(errors, ensure_ascii=True),
                ),
            )
            conn.commit()

    def list_runs(self, limit: int = 30) -> list[dict[str, Any]]:
        with self._conn() as conn:
            rows = conn.execute(
                """
                SELECT run_id, topic, created_at, critic_score, metrics_json
                FROM run_history
                ORDER BY created_at DESC
                LIMIT ?
                """,
                (limit,),
            ).fetchall()
        out: list[dict[str, Any]] = []
        for r in rows:
            metrics = json.loads(r["metrics_json"]) if r["metrics_json"] else {}
            out.append(
                {
                    "run_id": r["run_id"],
                    "topic": r["topic"],
                    "created_at": r["created_at"],
                    "critic_score": r["critic_score"],
                    "total_latency_s": metrics.get("total_latency_s"),
                }
            )
        return out

    def load_run(self, run_id: str) -> dict[str, Any] | None:
        with self._conn() as conn:
            row = conn.execute("SELECT * FROM run_history WHERE run_id = ?", (run_id,)).fetchone()
        if not row:
            return None
        return {
            "run_id": row["run_id"],
            "topic": row["topic"],
            "created_at": row["created_at"],
            "model_params": json.loads(row["model_params"] or "{}"),
            "sources": json.loads(row["sources_json"] or "[]"),
            "approved_source_ids": json.loads(row["approved_source_ids_json"] or "[]"),
            "search_results": row["search_results"] or "",
            "scraped_content": row["scraped_content"] or "",
            "report": row["report"] or "",
            "refined_report": row["refined_report"] or "",
            "feedback": row["feedback"] or "",
            "critic_score": row["critic_score"],
            "metrics": json.loads(row["metrics_json"] or "{}"),
            "errors": json.loads(row["errors_json"] or "{}"),
        }


def new_run_id() -> str:
    return uuid.uuid4().hex


def export_bundle(
    run_id: str,
    topic: str,
    report: str,
    feedback: str,
    metadata: dict[str, Any],
    output_dir: str = "exports",
) -> dict[str, str]:
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    ts = int(time.time())
    base = Path(output_dir) / f"{run_id}_{ts}"

    md_path = f"{base}.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(report)

    json_path = f"{base}.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(
            {
                "run_id": run_id,
                "topic": topic,
                "report": report,
                "feedback": feedback,
                "metadata": metadata,
            },
            f,
            ensure_ascii=True,
            indent=2,
        )

    docx_path = f"{base}.docx"
    doc = Document()
    doc.add_heading(topic, level=1)
    for paragraph in report.split("\n\n"):
        doc.add_paragraph(paragraph.strip())
    doc.add_page_break()
    doc.add_heading("Critic Feedback", level=2)
    doc.add_paragraph(feedback or "")
    doc.save(docx_path)

    pdf_path = f"{base}.pdf"
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    y = height - 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, topic[:95])
    y -= 28
    c.setFont("Helvetica", 10)
    text = c.beginText(40, y)
    text.setLeading(14)
    for line in (report + "\n\nCritic Feedback:\n" + (feedback or "")).splitlines():
        if y < 70:
            c.drawText(text)
            c.showPage()
            y = height - 40
            text = c.beginText(40, y)
            text.setLeading(14)
        text.textLine(line[:120])
        y -= 14
    c.drawText(text)
    c.save()

    return {
        "markdown": md_path,
        "json": json_path,
        "docx": docx_path,
        "pdf": pdf_path,
    }

